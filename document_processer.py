"""
Advanced Document Processor for Multi-Format Document Ingestion
Handles PDF, TXT, email, and other document types with OCR and table extraction
"""

import os
import re
import hashlib
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
import email
import json

# Document processing libraries
import fitz  # PyMuPDF
try:
    import pytesseract
    from PIL import Image
    import cv2
    import numpy as np
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ pytesseract not available. OCR functionality will be disabled.")
from pdf2image import convert_from_path
from docx import Document
from bs4 import BeautifulSoup
import requests
import pandas as pd
# Table extraction libraries
try:
    import tabula
    import camelot
    TABLE_EXTRACTION_AVAILABLE = True
except ImportError:
    TABLE_EXTRACTION_AVAILABLE = False
    print("⚠️ Table extraction libraries not available. Table extraction will be disabled.")

# LangChain for text splitting
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Represents a processed document chunk with metadata"""
    chunk_id: str
    content: str
    source_file: str
    file_type: str
    page_number: Optional[int] = None
    chunk_index: Optional[int] = None
    section_type: Optional[str] = None  # text, table, header, footer, etc.
    confidence_score: Optional[float] = None
    extracted_entities: Optional[Dict[str, Any]] = None
    table_data: Optional[Dict[str, Any]] = None
    embedding: Optional[List[float]] = None

class AdvancedDocumentProcessor:
    """Advanced document processor with multi-format support and table extraction"""
    
    def __init__(self, ocr_language='eng', chunk_size=400, chunk_overlap=50, max_chunks_to_consider=5):
        self.ocr_language = ocr_language
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_chunks_to_consider = max_chunks_to_consider
        
        # Initialize text splitter with semantic-aware sentence-based chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", "! ", "? ", "; ", ": ", " ", ""],
            keep_separator=True,  # Keep separators for better semantic boundaries
            length_function=len
        )
        
        # Supported file types
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.txt': self._process_txt,
            '.docx': self._process_docx,
            '.html': self._process_html,
            '.htm': self._process_html,
            '.eml': self._process_email,
            '.msg': self._process_email,
            '.csv': self._process_csv,
            '.json': self._process_json
        }
    
    def process_document(self, file_path: str, use_ocr: bool = False) -> List[DocumentChunk]:
        """Main entry point for document processing"""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            logger.info(f"Processing document: {file_path}")
            
            # Process based on file type
            processor_func = self.supported_extensions[file_extension]
            raw_content = processor_func(str(file_path), use_ocr)
            
            # Extract tables and structured content
            structured_content = self._extract_structured_content(raw_content, file_path)
            
            # Chunk the content
            chunks = self._chunk_content(structured_content, str(file_path), file_extension)
            
            logger.info(f"Successfully processed {len(chunks)} chunks from {file_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            logger.error("This might be due to missing dependencies or OCR issues.")
            logger.error("Please check:")
            logger.error("1. File exists and is readable")
            logger.error("2. Tesseract is installed (run: python fix_tesseract.py)")
            logger.error("3. All dependencies are installed (run: python install_nltk_version.py)")
            logger.error("4. Run debug_document_processing.py to diagnose issues")
            import traceback
            logger.error(f"Full error traceback: {traceback.format_exc()}")
            raise
    
    def _process_pdf(self, pdf_path: str, use_ocr: bool = False) -> Dict[str, Any]:
        """Process PDF with text extraction, OCR, and table detection"""
        try:
            doc = fitz.open(pdf_path)
            content = {
                'text': "",
                'tables': [],
                'images': [],
                'metadata': {}
            }
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text normally first
                page_text = page.get_text()
                
                # If no text or very little text, use OCR
                if use_ocr or len(page_text.strip()) < 50:
                    logger.info(f"Using OCR for page {page_num + 1}")
                    ocr_text = self._extract_text_with_ocr(page, page_num)
                    if ocr_text.strip():  # Only use OCR text if it's not empty
                        page_text = ocr_text
                    else:
                        logger.warning(f"OCR returned empty text for page {page_num + 1}, using normal extraction")
                
                content['text'] += f"\n\n--- Page {page_num + 1} ---\n{page_text}"
                
                # Extract tables from this page
                page_tables = self._extract_tables_from_page(page, page_num)
                content['tables'].extend(page_tables)
                
                # Extract images (for future OCR if needed)
                page_images = self._extract_images_from_page(page, page_num)
                content['images'].extend(page_images)
            
            doc.close()
            return content
            
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path}: {e}")
            raise
    
    def _extract_text_with_ocr(self, page, page_num: int) -> str:
        """Extract text using OCR with pytesseract"""
        if not OCR_AVAILABLE:
            logger.error("pytesseract not available. Please install it with: pip install pytesseract")
            return ""
        
        try:
            # Get page as image
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Simple OCR without complex preprocessing
            text = pytesseract.image_to_string(
                img, 
                lang=self.ocr_language,
                config='--psm 6'  # Assume uniform block of text
            )
            
            return text
            
        except Exception as e:
            logger.error(f"OCR error on page {page_num}: {e}")
            logger.error("This might be due to Tesseract not being installed or configured properly.")
            logger.error("Please run: python fix_tesseract.py")
            return ""
    
    def _extract_tables_from_page(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Extract tables from PDF page using basic methods"""
        tables = []
        
        try:
            if TABLE_EXTRACTION_AVAILABLE:
                # Method 1: Use camelot for table extraction
                try:
                    page_tables = camelot.read_pdf(
                        page.parent, 
                        pages=str(page_num + 1),
                        flavor='lattice'
                    )
                    
                    for table in page_tables:
                        if table.df.shape[0] > 1:  # Only include tables with data
                            tables.append({
                                'page': page_num + 1,
                                'data': table.df.to_dict('records'),
                                'accuracy': table.accuracy,
                                'whitespace': table.whitespace,
                                'method': 'camelot'
                            })
                except Exception as e:
                    logger.debug(f"Camelot table extraction failed: {e}")
                
                # Method 2: Use tabula for simpler tables
                try:
                    tabula_tables = tabula.read_pdf(
                        page.parent, 
                        pages=page_num + 1,
                        multiple_tables=True
                    )
                    
                    for i, table in enumerate(tabula_tables):
                        if not table.empty:
                            tables.append({
                                'page': page_num + 1,
                                'data': table.to_dict('records'),
                                'method': 'tabula'
                            })
                except Exception as e:
                    logger.debug(f"Tabula table extraction failed: {e}")
            
            # Method 3: Basic table extraction using text patterns (fallback)
            if not tables:
                tables = self._extract_tables_basic(page, page_num)
                
        except Exception as e:
            logger.error(f"Table extraction error on page {page_num}: {e}")
        
        return tables
    
    def _extract_images_from_page(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Extract images from PDF page"""
        images = []
        
        try:
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                pix = fitz.Pixmap(page.parent, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    images.append({
                        'page': page_num + 1,
                        'index': img_index,
                        'width': pix.width,
                        'height': pix.height,
                        'format': pix.colorspace.name
                    })
                
                pix = None  # Free memory
                
        except Exception as e:
            logger.error(f"Image extraction error on page {page_num}: {e}")
        
        return images
    
    def _extract_tables_basic(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Basic table extraction using text patterns when advanced libraries are not available"""
        tables = []
        
        try:
            # Get page text
            page_text = page.get_text()
            
            # Simple table detection using common patterns
            lines = page_text.split('\n')
            table_lines = []
            in_table = False
            
            for line in lines:
                # Check for table-like patterns (multiple columns separated by spaces/tabs)
                if len(line.strip()) > 0 and ('\t' in line or line.count('  ') >= 2):
                    if not in_table:
                        in_table = True
                        table_lines = []
                    table_lines.append(line)
                elif in_table:
                    # End of table detected
                    if table_lines:
                        # Convert to table format
                        table_data = []
                        for table_line in table_lines:
                            # Split by multiple spaces or tabs
                            columns = [col.strip() for col in table_line.split('\t') if col.strip()]
                            if not columns:
                                columns = [col.strip() for col in line.split('  ') if col.strip()]
                            if columns:
                                table_data.append(columns)
                        
                        if table_data and len(table_data) > 1:
                            tables.append({
                                'page': page_num + 1,
                                'data': table_data,
                                'method': 'basic_pattern'
                            })
                    
                    in_table = False
                    table_lines = []
            
            # Handle table at end of page
            if in_table and table_lines:
                table_data = []
                for table_line in table_lines:
                    columns = [col.strip() for col in table_line.split('\t') if col.strip()]
                    if not columns:
                        columns = [col.strip() for col in line.split('  ') if col.strip()]
                    if columns:
                        table_data.append(columns)
                
                if table_data and len(table_data) > 1:
                    tables.append({
                        'page': page_num + 1,
                        'data': table_data,
                        'method': 'basic_pattern'
                    })
                    
        except Exception as e:
            logger.error(f"Basic table extraction error on page {page_num}: {e}")
        
        return tables
    
    def _process_txt(self, txt_path: str, use_ocr: bool = False) -> Dict[str, Any]:
        """Process plain text files"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            return {
                'text': text,
                'tables': [],
                'images': [],
                'metadata': {}
            }
            
        except Exception as e:
            logger.error(f"Error processing TXT {txt_path}: {e}")
            raise
    
    def _process_docx(self, docx_path: str, use_ocr: bool = False) -> Dict[str, Any]:
        """Process Word documents"""
        try:
            doc = Document(docx_path)
            text = ""
            tables = []
            
            # Extract text from paragraphs - handle different versions of python-docx
            try:
                # Try the standard way first
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
            except AttributeError:
                # Fallback for different versions
                try:
                    # Try alternative attribute names
                    if hasattr(doc, 'paragraphs'):
                        for paragraph in doc.paragraphs:
                            text += paragraph.text + "\n"
                    elif hasattr(doc, 'content'):
                        text = doc.content
                    else:
                        # Last resort - try to extract text from the document structure
                        text = str(doc)
                except Exception as e:
                    logger.warning(f"Could not extract text from DOCX paragraphs: {e}")
                    text = "DOCX content could not be extracted"
            
            # Extract tables from Word document
            try:
                # Check if tables attribute exists
                if hasattr(doc, 'tables'):
                    for table in doc.tables:
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text for cell in row.cells]
                            table_data.append(row_data)
                        
                        if table_data:
                            tables.append({
                                'data': table_data,
                                'method': 'docx'
                            })
                else:
                    logger.warning("DOCX document has no 'tables' attribute - tables will not be extracted")
            except Exception as e:
                logger.warning(f"Could not extract tables from DOCX: {e}")
            
            return {
                'text': text,
                'tables': tables,
                'images': [],
                'metadata': {}
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX {docx_path}: {e}")
            raise
    
    def _process_html(self, html_path: str, use_ocr: bool = False) -> Dict[str, Any]:
        """Process HTML files"""
        try:
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text
            text = soup.get_text(separator='\n', strip=True)
            
            # Extract tables
            tables = []
            for table in soup.find_all('table'):
                table_data = []
                for row in table.find_all('tr'):
                    row_data = [cell.get_text(strip=True) for cell in row.find_all(['td', 'th'])]
                    if row_data:
                        table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        'data': table_data,
                        'method': 'html'
                    })
            
            return {
                'text': text,
                'tables': tables,
                'images': [],
                'metadata': {}
            }
            
        except Exception as e:
            logger.error(f"Error processing HTML {html_path}: {e}")
            raise
    
    def _process_email(self, email_path: str, use_ocr: bool = False) -> Dict[str, Any]:
        """Process email files (.eml, .msg)"""
        try:
            with open(email_path, 'r', encoding='utf-8') as f:
                email_content = f.read()
            
            # Parse email
            msg = email.message_from_string(email_content)
            
            # Extract email components
            subject = msg.get('Subject', '')
            sender = msg.get('From', '')
            recipient = msg.get('To', '')
            date = msg.get('Date', '')
            
            # Extract body
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body = part.get_payload(decode=True).decode()
                        break
            else:
                body = msg.get_payload(decode=True).decode()
            
            # Combine all text
            text = f"Subject: {subject}\nFrom: {sender}\nTo: {recipient}\nDate: {date}\n\n{body}"
            
            return {
                'text': text,
                'tables': [],
                'images': [],
                'metadata': {
                    'subject': subject,
                    'sender': sender,
                    'recipient': recipient,
                    'date': date
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing email {email_path}: {e}")
            raise
    
    def _process_csv(self, csv_path: str, use_ocr: bool = False) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            df = pd.read_csv(csv_path)
            
            # Convert to text representation
            text = df.to_string(index=False)
            
            # Store as table
            tables = [{
                'data': df.to_dict('records'),
                'method': 'csv'
            }]
            
            return {
                'text': text,
                'tables': tables,
                'images': [],
                'metadata': {
                    'columns': list(df.columns),
                    'rows': len(df)
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing CSV {csv_path}: {e}")
            raise
    
    def _process_json(self, json_path: str, use_ocr: bool = False) -> Dict[str, Any]:
        """Process JSON files"""
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert to text representation
            text = json.dumps(data, indent=2)
            
            return {
                'text': text,
                'tables': [],
                'images': [],
                'metadata': {
                    'type': 'json',
                    'keys': list(data.keys()) if isinstance(data, dict) else []
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing JSON {json_path}: {e}")
            raise
    
    def _extract_structured_content(self, content: Dict[str, Any], file_path: Path) -> List[Dict[str, Any]]:
        """Extract and structure content from processed document"""
        structured_content = []
        
        # Add main text content
        if content['text'].strip():
            structured_content.append({
                'type': 'text',
                'content': content['text'],
                'section_type': 'main_text'
            })
        
        # Add tables
        for i, table in enumerate(content['tables']):
            table_text = self._table_to_text(table['data'])
            structured_content.append({
                'type': 'table',
                'content': table_text,
                'section_type': 'table',
                'table_data': table,
                'table_index': i
            })
        
        # Add metadata as text
        if content.get('metadata'):
            metadata_text = json.dumps(content['metadata'], indent=2)
            structured_content.append({
                'type': 'text',
                'content': f"Document Metadata:\n{metadata_text}",
                'section_type': 'metadata'
            })
        
        return structured_content
    
    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """Convert table data to readable text"""
        if not table_data:
            return ""
        
        text_lines = []
        for row in table_data:
            text_lines.append(" | ".join(str(cell) for cell in row))
        
        return "\n".join(text_lines)
    
    def _chunk_content(self, structured_content: List[Dict[str, Any]], source_file: str, file_type: str) -> List[DocumentChunk]:
        """Chunk the structured content into DocumentChunk objects"""
        chunks = []
        chunk_index = 0
        
        for section in structured_content:
            try:
                # Use LangChain's text splitter for better semantic chunking
                docs = [Document(page_content=section['content'], metadata={"source": source_file})]
                split_docs = self.text_splitter.split_documents(docs)
                
                for i, doc in enumerate(split_docs):
                    chunk_id = f"chunk_{chunk_index+1}_{hashlib.md5(doc.page_content.encode()).hexdigest()[:8]}"
                    
                    chunk = DocumentChunk(
                        chunk_id=chunk_id,
                        content=doc.page_content.strip(),
                        source_file=source_file,
                        file_type=file_type,
                        chunk_index=chunk_index,
                        section_type=section.get('section_type', 'text'),
                        table_data=section.get('table_data'),
                        confidence_score=section.get('confidence_score', 1.0)
                    )
                    
                    chunks.append(chunk)
                    chunk_index += 1
                    
            except Exception as e:
                logger.error(f"Error chunking section: {e}")
                continue
        
        return chunks

# Example usage
if __name__ == "__main__":
    processor = AdvancedDocumentProcessor()
    
    # Test with a PDF file
    test_file = "sample.pdf"
    if os.path.exists(test_file):
        chunks = processor.process_document(test_file, use_ocr=False)
        print(f"Processed {len(chunks)} chunks from {test_file}")
        
        for i, chunk in enumerate(chunks[:3]):
            print(f"\nChunk {i+1}:")
            print(f"ID: {chunk.chunk_id}")
            print(f"Type: {chunk.section_type}")
            print(f"Content preview: {chunk.content[:100]}...")
    else:
        print(f"Test file {test_file} not found") 